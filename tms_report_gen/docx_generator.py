from pathlib import Path

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Inches

from tms_report_gen.models import TestCase
from tms_report_gen.utils import (
    format_value, 
    resolve_image_paths
)


def generate_docx(
    testcase: TestCase,
    output_path: str,
    media_root: str,
):

    doc = Document()

    # Heading - test case name
    doc.add_heading(testcase.name, level=1)

    # Test case attributes list
    doc.add_heading("Информация о тест-кейсе", level=2)

    add_attribute(doc, "ID", testcase.id)
    add_attribute(doc, "Title", testcase.title)
    add_attribute(doc, "Status", testcase.status)

    add_attribute(doc, "Tracked", testcase.is_tracked)
    add_attribute(doc, "Auto", testcase.is_auto)
    add_attribute(doc, "Individual", testcase.is_individ)
    add_attribute(doc, "Unique", testcase.is_unique)

    add_attribute(doc, "Priority", testcase.priority)
    add_attribute(doc, "Type", testcase.test_type)
    add_attribute(doc, "Complexity", testcase.complexity)

    # Steps section
    doc.add_heading("Шаги", level=2)

    for step in testcase.steps:

        title = (
            f"Step #{format_value(step.order)} "
            f"({step.step_type})"
        )

        doc.add_heading(title, level=3)

        parse_html(doc, step.action, media_root)

        if step.expected_result:

            p = doc.add_paragraph()

            run = p.add_run("Expected result:")
            run.bold = True

            parse_html(doc, step.expected_result, media_root)


    doc.save(output_path)


def add_attribute(
    doc: Document,
    name: str,
    value,
):

    p = doc.add_paragraph()

    key_run = p.add_run(f"{name}: ")
    key_run.bold = True

    p.add_run(format_value(value))


def parse_html(
    doc: Document,
    html: str,
    media_root: str,
):

    # Resolving image paths
    html = resolve_image_paths(
        html,
        media_root,
    )

    soup = BeautifulSoup(html, "html.parser")

    # Top-level elements handling
    for element in soup.contents:

        # Paragraphs
        if element.name == "p":

            p = doc.add_paragraph()

            # <b>, <i>, <code> and <br> tags handling (inside this paragraph)
            parse_inline_tags(p, element)

        
        # Code blocks
        elif element.name == "pre":

            code_text = element.get_text()

            p = doc.add_paragraph()

            run = p.add_run(code_text)

            run.font.name = "Courier New"
            run.font.size = Pt(10)


        # Lists
        elif element.name == "ul":
            parse_list(doc, element, ordered=False)

        elif element.name == "ol":
            parse_list(doc, element, ordered=True)


        # Tables
        elif element.name == "table":
            parse_table(doc, element)


        # Top-level images
        elif element.name == "img":

            p = doc.add_paragraph()

            add_image(
                p,
                element,
            )


        # Unhandled tags
        elif isinstance(element, str):

            text = element.strip()

            if text:
                doc.add_paragraph(text)


def parse_inline_tags(
    paragraph,
    parent,
):

    for child in parent.children:

        # Simple text
        if isinstance(child, str):

            paragraph.add_run(child)

        # Bold
        elif child.name == "b":

            run = paragraph.add_run(
                child.get_text()
            )

            run.bold = True

        # Italic
        elif child.name == "i":

            run = paragraph.add_run(
                child.get_text()
            )

            run.italic = True

        # Code
        elif child.name == "code":

            run = paragraph.add_run(
                child.get_text()
            )

            run.font.name = "Courier New"

        # Line break
        elif child.name == "br":

            paragraph.add_run("\n")

        # Hyperlink
        elif child.name == "a":

            add_link(
                paragraph,
                child
            )

        # Image
        elif child.name == "img":

            add_image(
                paragraph,
                child
            )
        
        # Unhandled tag
        else:

            paragraph.add_run(
                child.get_text()
            )


def parse_list(
    doc: Document,
    list_element,
    ordered: bool,
    level: int = 0,
):

    items = list_element.find_all(
        "li",
        recursive=False
    )

    # Enumerating through list items and adding them to the document
    for index, item in enumerate(items, start=1):

        p = doc.add_paragraph()

        # Adding indent for nested lists
        p.paragraph_format.left_indent = Inches(
            0.3 * level
        )

        # Defining prefix for ordered and unordered lists
        if ordered:
            prefix = f"{index}. "
        else:
            prefix = "• "

        p.add_run(prefix)

        # Parsing the content of the list item
        for child in item.children:

            # Text
            if isinstance(child, str):

                text = child.strip()

                if text:
                    p.add_run(text)

            # Nested lists
            elif child.name == "ul":

                parse_list(
                    doc,
                    child,
                    ordered=False,
                    level=level + 1,
                )

            elif child.name == "ol":

                parse_list(
                    doc,
                    child,
                    ordered=True,
                    level=level + 1,
                )

            # Hyperlink
            elif child.name == "a":

                add_link(
                    p,
                    child
                )

            # Other inline tags
            else:

                parse_inline_tags(
                    p,
                    child
                )


def parse_table(
    doc: Document,
    table_element,
):

    # Finding all rows in the table
    rows = table_element.find_all(
        "tr",
        recursive=False
    )

    if not rows:
        return

    first_row = rows[0]

    # Finding all columns in the first row to determine the number of columns in the table
    columns = first_row.find_all(
        ["td", "th"],
        recursive=False
    )

    # Creating the table in the document
    table = doc.add_table(
        rows=0,
        cols=len(columns)
    )

    table.style = "Table Grid"

    for row_element in rows:

        # Finding all cells in the current row
        cells = row_element.find_all(
            ["td", "th"],
            recursive=False
        )

        row = table.add_row().cells

        # Enumerating through the cells and adding them to the document table
        for index, cell in enumerate(cells):

            cell_paragraph = row[index].paragraphs[0]

            # Parsing the content of the cell
            parse_inline_tags(
                cell_paragraph,
                cell
            )


def add_link(
    paragraph,
    link_element,
):

    # Getting the text of the link
    text = link_element.get_text(
        strip=True
    )

    # Getting href attribute
    href = link_element.get(
        "href",
        ""
    )

    # Formatting the link text for display in the document
    if text and href:
        link_text = f"{text} ({href})"

    elif href:
        link_text = href

    else:
        link_text = text

    # Adding the link text to the paragraph with underline formatting
    run = paragraph.add_run(
        link_text
    )

    run.underline = True


def add_image(
    paragraph,
    img_element,
):

    src = img_element.get("src")

    # Handle invalid cases for the src attribute
    if not src:
        return

    if not src.startswith("file:///"):
        return

    try:

        # Resolving the image path from the src attribute
        image_path = Path(
            src.replace("file:///", "")
        )

        # Handle image absence with a placeholder
        if not image_path.exists():

            paragraph.add_run(
                f"[Image not found: {src}]"
            )

            return

        # Adding the image to the document
        run = paragraph.add_run()

        run.add_picture(
            str(image_path),
            width=Inches(5.5)
        )

    # Handle any exceptions so that the document generation doesn't fail
    except Exception:

        
        paragraph.add_run(
            f"[Failed to load image: {src}]"
        )